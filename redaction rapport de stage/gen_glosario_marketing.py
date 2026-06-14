# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x55, 0x55, 0x55)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, white=False, size=10.5, color=None, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        run.font.color.rgb = color

# --- encabezado de la sección de anexos ---
ph = doc.add_paragraph()
rh = ph.add_run("Anexos")
rh.bold = True
rh.font.size = Pt(16)
rh.font.color.rgb = NAVY
ph.paragraph_format.space_after = Pt(10)

# --- título del anexo ---
p = doc.add_paragraph()
r = p.add_run("Anexo 1. Glosario de términos de marketing y gestión")
r.bold = True
r.font.size = Pt(12.5)
r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(8)

pintro = doc.add_paragraph()
ri = pintro.add_run("Definiciones de los términos especializados de marketing, comercio y comunicación "
                    "digital empleados a lo largo del informe.")
ri.font.size = Pt(10.5)
pintro.paragraph_format.space_after = Pt(6)

terms = [
    # Conceptos de marketing y comercio (Kotler & Keller, 2012)
    ("B2B (business-to-business)",
     "Modelo comercial en el que una empresa vende sus productos o servicios a otras empresas o instituciones, y no al consumidor final."),
    ("Marketing de servicios",
     "Rama del marketing centrada en los servicios, que se distinguen de los bienes por ser intangibles, inseparables de quien los presta, variables y perecederos."),
    ("Intangibilidad",
     "Característica de un servicio que no puede verse, tocarse ni probarse antes de ser adquirido."),
    ("Centro de compra (buying center)",
     "Conjunto de personas que, dentro de una organización cliente, intervienen en la decisión de compra, con funciones distintas (iniciador, prescriptor, decisor, comprador y filtro)."),
    ("Estrategia push",
     "Estrategia en la que la empresa «empuja» activamente un producto hacia el cliente mediante prospección y promoción directa."),
    ("Estrategia pull",
     "Estrategia en la que es el cliente quien, atraído por la oferta, «tira» del producto y contacta él mismo a la empresa."),
    ("Segmentación de mercado",
     "División de un mercado en grupos de clientes con características o necesidades comunes, para adaptar la oferta y el mensaje a cada grupo."),
    ("Prospección comercial",
     "Conjunto de acciones destinadas a identificar y contactar a clientes potenciales."),
    ("Lead",
     "Contacto comercial potencialmente interesado, susceptible de convertirse en cliente."),
    ("Ciclo de venta",
     "Tiempo y etapas que transcurren entre el primer contacto con un cliente potencial y la firma del contrato."),
    ("MICE (Meetings, Incentives, Conferences and Exhibitions)",
     "Segmento del turismo de negocios que agrupa reuniones, viajes de incentivo, congresos y ferias profesionales."),
    # Marketing digital / e-mailing (Mailchimp, s. f.)
    ("Email marketing",
     "Técnica de marketing que consiste en enviar mensajes comerciales o informativos por correo electrónico a una lista de contactos."),
    ("Newsletter",
     "Boletín informativo enviado periódicamente por correo electrónico a una lista de suscriptores voluntarios."),
    ("Template (plantilla)",
     "Modelo de correo prediseñado que se reutiliza y se adapta para cada nueva campaña, garantizando una presentación uniforme."),
    ("Tasa de apertura (open rate)",
     "Porcentaje de correos entregados que fueron abiertos por los destinatarios."),
    ("Tasa de clics (click rate)",
     "Porcentaje de correos entregados en los que el destinatario hizo al menos un clic."),
    ("Tasa de rebote (bounce rate)",
     "Porcentaje de correos que no pudieron ser entregados, por direcciones no válidas u otros problemas técnicos."),
    ("Tasa de respuesta",
     "Porcentaje de contactos que responden de forma efectiva a una comunicación comercial."),
    ("Diseño UX (experiencia de usuario)",
     "Diseño orientado a que el usuario interactúe con un producto, una web o una interfaz de forma fácil, clara y satisfactoria."),
]

t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = 'Table Grid'
hc = t.rows[0].cells
set_cell(hc[0], "Término", bold=True, white=True, size=10.5); shade(hc[0], '1F3864')
set_cell(hc[1], "Definición", bold=True, white=True, size=10.5); shade(hc[1], '1F3864')
for term, defi in terms:
    cells = t.add_row().cells
    set_cell(cells[0], term, bold=True, size=10)
    set_cell(cells[1], defi, size=10)

widths = [Cm(4.8), Cm(12.0)]
for row in t.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w

# --- Nota (fuentes) ---
nota = doc.add_paragraph()
nota.paragraph_format.space_before = Pt(6)
nr = nota.add_run("Nota. Elaboración propia a partir de Kotler y Keller (2012) —conceptos de marketing y "
                  "comercio— y de la documentación de marketing digital de Mailchimp (s. f.) —métricas de "
                  "email marketing—.")
nr.italic = True
nr.font.size = Pt(10)

doc.save("Tabla 2 - Glosario marketing.docx")
print("Tabla 2 generada correctamente.")
