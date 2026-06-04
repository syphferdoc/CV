#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera 'Portada y agradecimientos.docx' conforme a la hoja de estilo LEA:
Times New Roman 12, interlineado 1,5, portada centrada + salto de página + agradecimientos justificados."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION

FONT = "Times New Roman"

doc = Document()

# Estilo base: Times New Roman 12, interlineado 1,5
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)


def para(text="", *, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=12,
         italic=False, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p


# ===================== PORTADA =====================
para("Université Paul-Valéry Montpellier 3", bold=True, size=14, space_before=12)
para("UFR 2 — Langues Étrangères Appliquées (LEA)", size=12)

para(space_before=36)
para("Ania Sarah Alloul", bold=True, size=13)
para("Licence 3 LEA — Anglais-Espagnol", size=12)
para("Parcours Management International", size=12)

para(space_before=48)
para("Informe de prácticas", size=13, italic=True)
para(space_before=12)
para("El flamenco en los mercados internacionales del espectáculo:",
     bold=True, size=16)
para("la experiencia de Flamenco Agency", bold=True, size=16)

para(space_before=54)
para("Empresa de acogida: Flamenco Agency", size=12)
para("Atanasio Barrón 20, local interior", size=12)
para("41003 Sevilla (España)", size=12)

para(space_before=18)
para("Tutora de prácticas: Sra. Alexis Stainow", size=12)
para("Período de prácticas: 15 de abril – 15 de junio de 2026", size=12)

para(space_before=54)
para("Año universitario 2025-2026", bold=True, size=12)

# ===================== AGRADECIMIENTOS (página nueva) =====================
doc.add_page_break()

para("Agradecimientos", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14,
     space_after=18)

agr = [
    "En primer lugar, quisiera expresar mi gratitud al conjunto del profesorado "
    "de la Universidad Paul-Valéry Montpellier 3, de quienes tanto he aprendido a "
    "lo largo de mi formación y cuya enseñanza ha sido la base de este trabajo.",

    "Mi agradecimiento más especial se dirige a mi tutora, la profesora Alya Ben "
    "Hamida, que fue también mi profesora de civilización española durante este "
    "curso. Sus clases me permitieron descubrir en profundidad la historia de "
    "España y, sobre todo, desarrollar una mirada crítica capaz de analizar cada "
    "aspecto de la España contemporánea sin caer en prejuicios. Esa transmisión "
    "del pensamiento analítico no solo me ha guiado en la redacción de este "
    "informe, sino también en mi vida cotidiana, pues me ha enseñado a reflexionar "
    "y a plantearme las preguntas adecuadas.",

    "Asimismo, deseo dar las gracias a todo el equipo de Flamenco Agency —y de "
    "manera especial al señor Jaime Trancoso, director de la agencia, y a la "
    "señora Alexis Stainow, mi tutora de prácticas— por haberme acogido con tanta "
    "amabilidad. Tuve la suerte de aprender en un entorno sano, benévolo y "
    "acogedor, que facilitó mi integración y me permitió crecer tanto en lo "
    "profesional como en lo personal. Hago extensivo este agradecimiento a mis "
    "compañeros becarios, con quienes compartí esta enriquecedora experiencia.",

    "Por último, deseo dar las gracias de todo corazón a mis padres, a mis "
    "hermanos y a Syphax, que siempre han creído en mí y me han apoyado en la "
    "reanudación de mis estudios tras haber obtenido mi máster en Argelia. Su "
    "presencia constante me dio fuerzas para sobreponerme a las dificultades que "
    "viví en Francia, al empezar de nuevo desde cero para alcanzar mi proyecto "
    "profesional.",
]

for t in agr:
    para(t, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=12)

out = "Portada y agradecimientos.docx"
doc.save(out)
print("OK ->", out)
