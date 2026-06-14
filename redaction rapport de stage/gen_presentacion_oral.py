# -*- coding: utf-8 -*-
"""Genera el documento Word de la presentacion oral (soutenance) + preguntas/respuestas."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- estilos base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11.5)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25

GREY = RGBColor(0x70, 0x70, 0x70)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x6B, 0x2E)


def h_title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(4)
    return p


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def timing(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(6)
    return p


def body(text):
    return doc.add_paragraph(text)


def optional(text):
    """Parrafo que se puede recortar si falta tiempo (en gris)."""
    p = doc.add_paragraph()
    r = p.add_run("[opcional, si te sobra tiempo] ")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREEN
    r2 = p.add_run(text)
    return p


def qa(q, a):
    p = doc.add_paragraph()
    r = p.add_run("P. " + q)
    r.bold = True
    r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(2)
    pa = doc.add_paragraph()
    ra = pa.add_run("R. ")
    ra.bold = True
    pa.add_run(a)
    pa.paragraph_format.space_after = Pt(10)


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("—" * 30)
    r.font.color.rgb = GREY


# ============================ PORTADA ============================
h_title("Presentación oral – Soutenance de prácticas")
sub = doc.add_paragraph()
r = sub.add_run("Ania Sarah Alloul · Licencia 3 LEA – Management Internacional · "
                "Universidad Paul-Valéry Montpellier 3")
r.font.size = Pt(11)
r.font.color.rgb = GREY
sub2 = doc.add_paragraph()
r = sub2.add_run("International Marketing Intern – Flamenco Agency (Sevilla) · "
                 "15 de abril – 15 de junio")
r.font.size = Pt(11)
r.font.color.rgb = GREY

note = doc.add_paragraph()
r = note.add_run("Exposición prevista: 10–15 minutos, en español. Las indicaciones de tiempo en "
                 "gris y las marcas «[opcional]» te ayudan a ajustar la duración el día de la "
                 "soutenance. Después vendrán unos 20 minutos de preguntas: encontrarás una batería de "
                 "preguntas probables con respuestas modelo en la última parte del documento.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY
hr()

# ============================ INTRODUCCIÓN ============================
h1("INTRODUCCIÓN – Presentar el marco y el sujeto")
timing("≈ 2 min 30")

body("Buenos días. Me llamo Ania Sarah Alloul y soy estudiante de tercer año de la Licencia LEA "
     "—Lenguas Extranjeras Aplicadas al management internacional— en la Universidad "
     "Paul-Valéry Montpellier 3. Durante dos meses, del 15 de abril al 15 de junio, realicé mis "
     "prácticas en Sevilla, en la empresa Flamenco Agency, como International Marketing Intern.")

body("Antes de entrar en el detalle, quisiera presentar brevemente la empresa y explicar por qué "
     "elegí estas prácticas. Flamenco Agency es una microempresa con sede en el centro de Sevilla, "
     "fundada en 2007 por Jaime Trancoso, doctor en musicología y especialista en flamenco, y "
     "dirigida junto a su esposa, Alexis Stainow, agente artística, que fue mi tutora de prácticas. "
     "Su actividad consiste en producir y distribuir espectáculos de flamenco y danza española en "
     "los mercados internacionales: desde su creación, la agencia ha trabajado en más de 25 países "
     "y en cuatro continentes, con festivales y teatros de referencia.")

body("Elegí estas prácticas porque correspondían exactamente a mi formación: la empresa buscaba un "
     "perfil con competencias en lenguas extranjeras, gestión y comunicación internacional. "
     "Encontré el contacto gracias al Sevilla Congress and Convention Bureau, que me facilitó una "
     "lista de empresas culturales, y envié una candidatura espontánea que fue aceptada.")

body("En cuanto al servicio y a mi papel dentro del equipo: la estructura permanente la formaban "
     "únicamente Jaime y Alexis, que trabajaban con cuatro becarios internacionales. Yo formaba "
     "parte de ese equipo de becarios, encargado del marketing internacional, es decir, de tres "
     "grandes ámbitos: la prospección comercial, la comunicación profesional multilingüe y el "
     "marketing digital.")

body("A lo largo de mis prácticas, una pregunta guió mi trabajo y mi informe: ¿cómo logra Flamenco "
     "Agency transformar el flamenco —entre folklore turístico y patrimonio cultural vivo— "
     "en una oferta artística competitiva en los mercados internacionales del espectáculo? Para "
     "responder, he organizado esta presentación en tres ejes: primero, el impacto de lo digital en "
     "las comunicaciones; segundo, un análisis entre mis expectativas y la realidad de la misión; y "
     "por último, el impacto de la inmersión cultural sevillana en mi percepción del flamenco.")

# ============================ EJE 1 ============================
h1("EJE 1 – El impacto de lo digital en las comunicaciones")
timing("≈ 4 min  ·  cubre: servicio y misiones, objetivos, dificultades y resultados con ejemplos")

body("Mi primer eje es el impacto de lo digital, porque toda la comunicación de la agencia con sus "
     "clientes se realiza, esencialmente, por medios digitales. Y para entenderlo hay que precisar "
     "algo fundamental: Flamenco Agency funciona según una lógica B2B, business-to-business. No "
     "vende sus espectáculos al gran público, sino a profesionales que deciden la programación: "
     "directores artísticos de festivales, programadores de teatros de ópera, responsables de "
     "temporadas culturales. Todo este contacto se desarrollaba casi exclusivamente a distancia, "
     "por correo electrónico y, en una fase más avanzada, por teléfono. La oficina de Sevilla era "
     "un centro de prospección y de gestión, no un punto de venta. Por eso las herramientas "
     "digitales eran el corazón de mi trabajo.")

body("Mi primera misión fue la prospección comercial. El objetivo era identificar y contactar "
     "festivales y teatros susceptibles de programar nuestros espectáculos para las temporadas 2027 "
     "y 2028. Trabajábamos sobre todo con una base de datos del Festival de Aviñón, con más de "
     "2.500 contactos, que había que clasificar y filtrar para conservar solo los leads pertinentes "
     "—directores artísticos y programadores— y eliminar el resto. Cuando faltaba un "
     "correo, lo buscaba en LinkedIn o en las webs oficiales de los festivales. Esta tarea, la "
     "gestión de esa base de datos, fue quizás la más formadora: me enseñó de manera muy concreta "
     "qué es la segmentación de un mercado.")

body("Mi segunda misión era redactar las propuestas comerciales. Cada correo presentaba tres "
     "espectáculos del catálogo, seleccionados según el perfil del destinatario: el tipo de "
     "festival, el país, el tamaño del escenario. Las redactaba principalmente en inglés, y en "
     "español para España y América Latina. Durante los dos meses envié más de 200 propuestas "
     "personalizadas.")

body("Mi tercera misión era el email marketing con Mailchimp. Aquí ya no se trataba de correos "
     "individuales, sino de campañas visuales dirigidas a toda la red de contactos para anunciar "
     "las giras confirmadas. Diseñé varias campañas, por ejemplo para Sangre Gitana en Praga o "
     "Argentine Tango en Letonia.")

body("Quiero dar un ejemplo concreto de resultados, porque ilustra muy bien el valor de lo digital. "
     "La campaña de Sangre Gitana fue la única que se envió efectivamente durante mis prácticas, "
     "así que pude consultar todos sus datos. Se dirigió a unos 720 contactos, con una tasa de "
     "apertura del 11,6 % y una tasa de clics del 0,9 %. Son cifras modestas en apariencia, pero "
     "hay que interpretarlas en el contexto de una prospección B2B « en frío »: los "
     "destinatarios no se habían suscrito voluntariamente. El dato más interesante fue otro: el "
     "enlace más clicado, con diferencia, fue el vídeo de YouTube del espectáculo, muy por delante "
     "de cualquier otro. Esto confirma algo esencial: un espectáculo es un producto intangible, que "
     "el cliente no puede ver antes de programarlo, y el vídeo es la herramienta digital que mejor "
     "lo hace « tangible ».")

optional("Tuve también una misión más personal: la traducción de la web al árabe, una lengua que "
         "conozco desde la infancia, dentro de la estrategia de expansión de la agencia hacia "
         "Oriente Medio y el Magreb.")

body("La principal dificultad de este eje fue técnica: ninguna asignatura me había preparado para "
     "usar profesionalmente Mailchimp, Canva o las funciones avanzadas de Excel. Tuve que "
     "aprenderlas directamente en contexto. Pero precisamente eso fue un resultado importante: salí "
     "de las prácticas dominando un conjunto de herramientas digitales profesionales que antes no "
     "conocía.")

# ============================ EJE 2 ============================
h1("EJE 2 – Análisis « expectativas frente a realidad » de mi misión")
timing("≈ 3 min 30  ·  cubre: objetivos de las misiones, dificultades, capacidad de análisis")

body("Mi segundo eje es un análisis más reflexivo: la diferencia entre lo que yo esperaba de la "
     "misión y la realidad que descubrí sobre el terreno. Comprender esa diferencia fue, "
     "probablemente, mi aprendizaje más importante.")

body("Cuando empecé, mi expectativa era bastante intuitiva: pensaba que el trabajo comercial "
     "consistía en enviar buenas propuestas y obtener respuestas, más o menos rápido. La realidad "
     "del marketing B2B en el sector cultural resultó ser muy distinta.")

body("Primera diferencia: el ritmo. Yo esperaba resultados relativamente rápidos. En realidad, el "
     "ciclo de venta es muy largo. Según me explicó el señor Trancoso, el tiempo medio entre el "
     "primer contacto y la firma de un contrato es de unos seis meses, y a veces los teatros "
     "deciden con uno o dos años de antelación. Hay excepciones —la República Checa llegó a "
     "contratar un espectáculo apenas un mes antes—, pero la norma es la paciencia.")

body("Segunda diferencia: la tasa de éxito. Esto me sorprendió mucho. La tasa de respuesta positiva "
     "en este sector ronda el 1 % de los contactos. Es decir, hay que aceptar que la inmensa "
     "mayoría de las propuestas no recibirán respuesta, y que el trabajo se mide en volumen y en "
     "constancia, no en resultados inmediatos. Debo ser honesta: durante mis dos meses, las "
     "campañas de prospección no generaron contratos firmados. Pero aprendí que ese no es el único "
     "objetivo: cada campaña mantiene la visibilidad de la empresa ante los decisores a largo plazo.")

body("Tercera diferencia: el cliente no es una sola persona. En el B2B interviene lo que se llama "
     "un « centro de compra »: un iniciador, un influenciador, un decisor, un comprador y, "
     "a veces, un « gatekeeper » que filtra el acceso. Aprendí a identificar a quién debía "
     "dirigirme, porque un mensaje excelente enviado a la persona equivocada no produce ningún "
     "resultado.")

body("A esta realidad profesional se sumaron dificultades lingüísticas que tampoco esperaba. Mi "
     "nivel de inglés es C1, así que el desafío no era la lengua en sí, sino el vocabulario "
     "especializado del sector: expresiones como « artistic line », « technical "
     "rider » o « booking conditions », que no había trabajado en la universidad. En "
     "español me defiendo bien, pero descubrí una dificultad inesperada: el habla andaluza. Las dos "
     "becarias sevillanas hablaban muy rápido y con muchas expresiones locales —me explicaron "
     "palabras como « mu » por « muy », o « la caló » por « el "
     "calor »— que no existen en el español académico. Fue, a la vez, un hándicap al "
     "principio y una lección sobre la distancia entre la lengua de la universidad y la lengua viva.")

body("¿Qué balance saco de este eje? Que mis expectativas iniciales eran las de una estudiante, y "
     "que la realidad me enseñó cómo funciona de verdad el comercio internacional: con ciclos "
     "largos, tasas de éxito bajas, un proceso de decisión complejo y la necesidad constante de "
     "adaptar el mensaje. Lejos de decepcionarme, esa diferencia entre expectativa y realidad es "
     "exactamente lo que vine a aprender.")

# ============================ EJE 3 ============================
h1("EJE 3 – El impacto de la inmersión cultural sevillana en mi percepción del flamenco")
timing("≈ 3 min  ·  incluye el ejemplo vivido de la Feria de Abril")

body("Mi tercer eje es más personal y cultural: cómo la inmersión en Sevilla transformó mi propia "
     "percepción del flamenco, que está en el centro de la problemática de mi informe.")

body("Llegué a Sevilla una semana antes de empezar a trabajar, justo para la Feria de Abril, una "
     "de las fiestas más importantes de Andalucía. Y esa experiencia cambió mi mirada antes incluso "
     "de pisar la oficina. Cuando asistí a la Feria, comprendí que el flamenco forma parte "
     "integrante de la vida de los sevillanos. No era solo turismo ni folklore, como a veces lo "
     "presenta la industria turística: era una verdadera manera de vivir, un modo de reunir a la "
     "gente, en lo familiar y en lo amistoso. Y, sobre todo, contaba historias: el flamenco puede "
     "ser triste, melancólico, alegre, puede hablar de amor. Tiene muchísimas historias que contar, "
     "y no es únicamente folklore.")

body("Esta vivencia personal conecta directamente con la pregunta de mi informe. El flamenco se "
     "encuentra en una situación ambivalente: a veces se percibe como un simple folklore turístico, "
     "y a veces circula como un patrimonio cultural vivo en los grandes circuitos internacionales. "
     "Lo que descubrí en la Feria —el flamenco como arte vivo y portador de emociones— es "
     "justamente lo que la agencia intenta vender en los mercados profesionales.")

body("Y así entendí mejor la estrategia de Flamenco Agency. La agencia no parte de cero para "
     "convencer a un programador, porque el flamenco ya cuenta con un fuerte capital simbólico como "
     "arte. Por ejemplo, descubrí la figura del bailaor Vicente Escudero, que ya en los años veinte "
     "triunfó en París, se relacionó con artistas de vanguardia como Picasso o Cocteau, y demostró "
     "que el flamenco tenía una verdadera dignidad teatral. Por eso, sobre todo en Francia, el "
     "flamenco está asociado desde hace casi un siglo al arte, y no al simple espectáculo turístico.")

body("La agencia se apoya en esa legitimidad y, además, la refuerza: elige cuidadosamente a sus "
     "interlocutores —grandes teatros y festivales que ya conocen el flamenco—, participa "
     "en ferias profesionales como WOMEX u Opera Europa, donde el contacto cara a cara consolida la "
     "confianza, y pertenece a redes de referencia como la ISPA. Un ejemplo significativo es el "
     "Teatro Campoamor de Oviedo: fue el propio teatro quien contactó a la agencia, y no al revés.")

body("En resumen, este eje me enseñó algo que ningún manual me habría enseñado: que detrás de una "
     "estrategia de marketing hay una realidad cultural viva, y que vender flamenco en el "
     "extranjero es, en el fondo, transmitir esas historias que descubrí en la Feria de Abril.")

# ============================ CONCLUSIÓN ============================
h1("CONCLUSIÓN – Reflexión personal: lo que me ha aportado")
timing("≈ 2 min")

body("Para concluir, quisiera compartir una breve reflexión personal sobre lo que esta experiencia "
     "me ha aportado, en tres planos.")

body("En el plano profesional, estas prácticas me han dado una base sólida en el funcionamiento del "
     "B2B internacional: la prospección, la segmentación de un mercado, el ciclo de venta y la "
     "paciencia que exige. En el plano lingüístico, confirmaron que la combinación de lenguas de la "
     "formación LEA responde a una necesidad real: pasar del inglés al español, y movilizar el "
     "árabe o el francés según el interlocutor, fue una ventaja constante. Y en el plano personal, "
     "vivir en una microempresa y en un entorno cultural distinto reforzó mi autonomía, mi "
     "iniciativa y mi confianza.")

body("Mi proyecto profesional se orienta hacia el comercio internacional, y más concretamente hacia "
     "la importación y la exportación. Aunque no deseo continuar mi carrera en el sector del arte y "
     "la cultura, estas prácticas me han permitido descubrir una faceta inesperada: la de la "
     "exportación cultural. He comprendido que un bien cultural se comercializa en el extranjero "
     "según una lógica comparable a la de cualquier otro producto —con sus mercados, su "
     "prospección y sus márgenes—, pero con la particularidad de ser un servicio intangible y "
     "de un fuerte valor simbólico.")

body("Por eso puedo decir que las prácticas en Flamenco Agency no solo completaron mi formación de "
     "LEA, sino que trazaron un puente directo hacia el ámbito profesional al que aspiro: el "
     "comercio internacional, en un contexto internacional y, quizás, en el extranjero. Muchas "
     "gracias por su atención; quedo a su disposición para responder a sus preguntas.")

# ============================ PREGUNTAS Y RESPUESTAS ============================
doc.add_page_break()
h_title("Posibles preguntas del jurado y respuestas modelo")
note = doc.add_paragraph()
r = note.add_run("Los « 20 minutos de intercambio » evalúan tu profesionalidad, tu "
                 "capacidad de analizar la experiencia (lo aprendido y las dificultades), la lengua "
                 "y tu capacidad de respuesta. Aquí tienes preguntas probables agrupadas por tema, "
                 "con respuestas que puedes adaptar con tus propias palabras.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

h1("Sobre la empresa y la elección del stage")
qa("¿Por qué elegiste Flamenco Agency y no otra empresa?",
   "Buscaba unas prácticas en Sevilla que me permitieran aplicar mi formación LEA en un contexto "
   "internacional real. El Sevilla Congress and Convention Bureau me facilitó contactos del sector "
   "cultural y, entre ellos, Flamenco Agency buscaba justamente un perfil plurilingüe con "
   "competencias en gestión y comunicación internacional. Me atrajo además trabajar en una "
   "microempresa con proyección en más de 25 países, porque ahí un becario asume responsabilidades "
   "reales desde el principio.")
qa("¿En qué consistía exactamente la actividad de la empresa?",
   "Flamenco Agency produce y distribuye espectáculos de flamenco y danza española a escala "
   "internacional. No vende al público, sino a profesionales —teatros, óperas y festivales—, "
   "según una lógica B2B. Distribuye sobre todo producciones de dos compañías, el Barcelona "
   "Flamenco Ballet y el Granada Flamenco Ballet, además de formatos más pequeños y de eventos "
   "corporativos.")

h1("Sobre las misiones, los logros y las dificultades")
qa("¿Cuál fue tu mayor dificultad y cómo la superaste?",
   "Tuve dos tipos de dificultad. La técnica: no sabía usar profesionalmente Mailchimp, Canva ni "
   "las funciones avanzadas de Excel, y las aprendí sobre el terreno, observando, practicando y "
   "apoyándome en las plantillas de la agencia. Y la lingüística: el vocabulario especializado del "
   "sector y el habla andaluza, muy rápida y con expresiones locales. Lo superé con la exposición "
   "diaria y pidiendo que me explicaran lo que no entendía; al final del período redactaba y "
   "comprendía con mucha más soltura.")
qa("Las campañas no generaron contratos firmados. ¿Consideras entonces que tu trabajo fue útil? "
   "¿Cómo se mide el éxito?",
   "Sí, fue útil, porque en este sector el éxito no se mide en contratos inmediatos. El ciclo de "
   "venta dura una media de seis meses, a veces uno o dos años, y la tasa de respuesta positiva "
   "ronda el 1 %. El objetivo de la prospección es mantener la visibilidad de la agencia ante los "
   "decisores y construir una relación a largo plazo. Mi trabajo —limpiar las bases de datos, "
   "enviar más de 200 propuestas y preparar campañas— alimenta ese proceso, cuyos frutos suelen "
   "llegar meses después.")
qa("¿Qué competencia consideras que más desarrollaste?",
   "Diría que la autonomía y la capacidad de adaptar el mensaje al interlocutor. Aprendí a analizar "
   "el perfil de un festival o un teatro y a elegir qué espectáculos proponer y con qué tono. Y, en "
   "el plano técnico, gané un dominio real de las herramientas de marketing digital y de la gestión "
   "de bases de datos.")
qa("Has hablado de la base de datos de Aviñón. ¿Qué aprendiste con ella concretamente?",
   "Fue la tarea más formadora. Tenía más de 2.500 contactos muy variados —técnicos, medios, "
   "artistas— y mi trabajo era filtrarlos para conservar solo a los decisores: directores "
   "artísticos y programadores. Eso me enseñó, de manera muy práctica, lo que es la segmentación de "
   "un mercado, y también la importancia del rigor, porque una dirección errónea hace inútil todo el "
   "esfuerzo.")

h1("Sobre los conceptos de marketing (capacidad de análisis)")
qa("¿Qué es el marketing B2B y en qué se diferencia del B2C?",
   "El B2B, business-to-business, es la venta entre empresas o profesionales; el B2C es la venta al "
   "consumidor final. En Flamenco Agency el cliente no es el espectador, sino el profesional que "
   "programa: un director artístico o un programador. Eso cambia todo: la decisión es más racional, "
   "el ciclo es más largo, intervienen varias personas y la relación se construye a largo plazo.")
qa("¿Qué es un « centro de compra »?",
   "Es el conjunto de personas que intervienen en una decisión de compra dentro de una institución. "
   "Suele haber un iniciador, que detecta la necesidad; un influenciador, que orienta; un decisor, "
   "que aprueba; un comprador, que negocia; y a veces un « gatekeeper », que filtra el "
   "acceso a las personas clave. Saber a quién dirigirse es decisivo, porque un buen mensaje a la "
   "persona equivocada no sirve de nada.")
qa("Explicaste la diferencia entre estrategia « push » y « pull ». ¿Puedes dar "
   "un ejemplo?",
   "La estrategia push es la oferta activa: cinco producciones —como Carmen o Luxuria— se "
   "proponen sistemáticamente en las campañas porque dejan más margen. La estrategia pull es lo "
   "contrario: otras producciones solo figuran en la web y son los propios programadores quienes las "
   "descubren y contactan a la agencia. El caso del Teatro Campoamor de Oviedo es un buen ejemplo de "
   "pull: fue el teatro el que dio el primer paso.")

h1("Sobre la lengua y el perfil LEA")
qa("¿Qué aportó concretamente tu perfil plurilingüe a la empresa?",
   "Trabajé en cuatro lenguas: inglés y español como lenguas principales, y árabe y francés de "
   "manera puntual. El inglés y el español me permitían recibir cualquier tarea sin restricción. "
   "Pero el verdadero valor diferencial fue el árabe: traduje la web de la agencia al árabe, dentro "
   "de su estrategia de expansión hacia Oriente Medio y el Magreb, algo que ningún otro becario "
   "podía hacer. Eso demuestra que una competencia personal puede convertirse en un activo "
   "profesional.")
qa("¿Qué laguna detectaste en tu formación LEA?",
   "Sobre todo dos. La formación no me preparó para las herramientas profesionales —Mailchimp, "
   "Canva, Excel avanzado— ni para el vocabulario muy especializado del sector. Pero no lo veo "
   "como un fallo de la licencia: es una formación generalista, y precisamente las prácticas sirven "
   "para tender ese puente entre la universidad y el mundo profesional.")

h1("Sobre el flamenco, la cultura y el proyecto profesional")
qa("Según tu experiencia, ¿el flamenco es folklore turístico o patrimonio cultural vivo?",
   "Es las dos cosas a la vez, y ahí está toda la cuestión de mi informe. Para una parte del "
   "turismo sigue siendo un folklore; pero, como descubrí en la Feria de Abril, es ante todo un "
   "patrimonio cultural vivo, una manera de vivir que transmite emociones e historias. La estrategia "
   "de Flamenco Agency consiste justamente en desplazar la percepción del folklore hacia el arte, "
   "dirigiéndose a programadores que ya lo valoran como tal.")
qa("¿La Feria de Abril cambió algo en tu trabajo concreto?",
   "Cambió mi comprensión del producto. Después de la Feria entendí emocionalmente lo que vendía: "
   "no un cliché turístico, sino un arte vivo. Eso me ayudó a redactar propuestas más convincentes y "
   "a entender por qué el vídeo es tan importante: hay que transmitir esa fuerza, ese « duende "
   "», a un cliente que no puede ver el espectáculo de antemano.")
qa("Si tu proyecto es el comercio internacional, ¿qué sentido tiene un stage en el sector cultural?",
   "Mucho, porque descubrí que un bien cultural se exporta según una lógica comparable a la de "
   "cualquier otro producto: mercados jerarquizados, prospección, ciclo de venta y márgenes. "
   "Aprendí los fundamentos del B2B internacional, que son exactamente los que necesito para la "
   "importación y la exportación. La cultura fue el contenido; el comercio internacional, la "
   "competencia transferible.")
qa("Si volvieras a empezar, ¿qué harías de otra manera?",
   "Dedicaría desde el primer día más tiempo a estudiar el vocabulario especializado del sector, en "
   "inglés y en español, porque al principio me frenó. Y propondría medir mejor el seguimiento de "
   "las propuestas, para analizar con más datos qué tipo de mensaje funciona mejor según el país.")

h1("Pregunta de cierre habitual")
qa("En una frase, ¿qué te ha aportado este stage?",
   "Me ha dado una base real en el comercio internacional B2B y la confirmación de que quiero "
   "construir mi carrera en un contexto internacional, además de demostrarme que mi perfil "
   "plurilingüe es un verdadero valor profesional.")

doc.save("Presentación oral - soutenance.docx")
print("OK -> Presentación oral - soutenance.docx")
