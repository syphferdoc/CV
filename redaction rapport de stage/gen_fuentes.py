# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- estilo base ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x55, 0x55, 0x55)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=10, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        run.font.color.rgb = color

def header_row(table, labels):
    cells = table.rows[0].cells
    for c, lab in zip(cells, labels):
        set_cell(c, lab, bold=True, white=True, size=10)
        shade(c, '1F3864')

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(6)

# ============================== TITRE ==============================
title = doc.add_paragraph()
tr = title.add_run('Fuentes y citas del informe de prácticas')
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = NAVY
sub = doc.add_paragraph()
sr = sub.add_run('Flamenco Agency · Licence 3 LEA · Université Paul-Valéry Montpellier 3')
sr.font.size = Pt(10.5)
sr.font.color.rgb = GREY
doc.add_paragraph()

note("Document de travail (FR) pour éviter le plagiat : pour chaque élément qui ne vient pas de toi "
     "(donnée chiffrée, notion théorique, fait historique, définition linguistique), tu trouves ici la "
     "phrase exacte à insérer dans ton rapport et la source correspondante. À coller : la section « Fuentes » "
     "à la fin du rapport ; et à reporter dans le texte les mentions de la colonne « Cómo citarlo ».")

# ============================== 1. COMMENT CITER ==============================
h1("1. Comment ça marche (rappel APA)")
b1 = doc.add_paragraph(style='List Bullet')
b1.add_run("Dans le texte : ").bold = True
b1.add_run("on met (Apellido, año) entre parenthèses, ou « Según Apellido (año)… ». "
           "Cela signale que l'idée/le chiffre vient d'une source.")
b2 = doc.add_paragraph(style='List Bullet')
b2.add_run("À la fin : ").bold = True
b2.add_run("chaque source citée doit figurer dans la liste « Fuentes » (section 4), par ordre alphabétique.")
b3 = doc.add_paragraph(style='List Bullet')
b3.add_run("Ton vécu, tes chiffres de la campagne Sangre Gitana, ce que t'a dit M. Trancoso, "
           "le glossaire (« elaboración propia ») = ").bold = False
b3.add_run("pas besoin de source").bold = True
b3.add_run(" (c'est toi / source primaire).")

# ============================== 2. TABLEAU PRINCIPAL ==============================
h1("2. Qué señalar y cómo citarlo en el texto")
note("Colonne 3 = le texte exact à mettre dans ton rapport (la citation est en gras).")
note("À retenir : les NOMS PROPRES connus (Paco de Lucía, Carmen Amaya, Antonio Gades, WOMEX, "
     "Teatro Campoamor…) = conocimiento común → tu peux les nommer SANS source. On ne cite que pour "
     "appuyer une affirmation, un chiffre, une définition ou une idée empruntée.")

rows = [
    # Apartado, Elemento (no es tuyo), Cómo citarlo en el texto, Fuente
    ("Intro",
     "Empleo cultural en España (+700 000 personas)",
     "«…este sector emplea en España a más de 700.000 personas (Ministerio de Cultura, 2024).»",
     "Ministerio de Cultura (2024). Anuario de estadísticas culturales"),

    ("Intro",
     "Turismo cultural (motivación de los turistas)",
     "«Según datos de Turespaña, cerca de uno de cada cinco viajes de turistas internacionales tiene una motivación cultural (Turespaña, 2023).»",
     "Turespaña (2023)"),

    ("Intro",
     "Visitantes de Sevilla al año (más de 3 millones) — VERIFICADO ✓",
     "«Sevilla, con más de 3 millones de visitantes al año (Instituto de Estadística y Cartografía de Andalucía, 2024)…»  [Dato confirmado: 3.018.060 viajeros alojados en 2023.]",
     "IECA / INE – Encuesta de Ocupación Hotelera (2024)"),

    ("Intro",
     "Ranking de congresos (ICCA)",
     "«España es el tercer país del mundo en número de congresos internacionales, y Barcelona y Madrid figuran entre las diez ciudades más activas (ICCA, 2024).»",
     "ICCA (2024). ICCA Statistics Report 2023"),

    ("Intro",
     "Eventos profesionales en Sevilla (MICE)",
     "«…Sevilla figura entre las ciudades españolas más activas en eventos profesionales (Spain Convention Bureau, 2023).»",
     "Spain Convention Bureau (2023)"),

    ("Intro",
     "Flamenco = Patrimonio UNESCO 2010",
     "«…el flamenco fue reconocido como Patrimonio Cultural Inmaterial de la Humanidad por la UNESCO en 2010 (UNESCO, 2010).»",
     "UNESCO (2010)"),

    ("Intro",
     "Orígenes del flamenco (andaluzas, gitanas, norteafricanas)",
     "«Sus raíces vienen de una mezcla de influencias andaluzas, gitanas y norteafricanas (Junta de Andalucía, s. f.).»",
     "Junta de Andalucía – Instituto Andaluz del Flamenco (s. f.)"),

    ("Intro / III.3",
     "Carmen Amaya, Antonio Gades, Paco de Lucía (difusión internacional del flamenco, mediados del s. XX)",
     "Nombrarlos NO necesita cita (conocimiento común). La cita respalda la AFIRMACIÓN: «…artistas como Carmen Amaya, Antonio Gades o Paco de Lucía hicieron conocer el flamenco en Europa, América y Asia (Junta de Andalucía, s. f.).»",
     "Junta de Andalucía – Instituto Andaluz del Flamenco (s. f.)"),

    ("Intro / III.3",
     "Datos de la empresa, sus referencias, clientes y membresías (Royal Opera House Muscat, HP, ISPA, EFA, WOMEX…)",
     "FUENTE PRIMARIA: es tu experiencia y los materiales internos de la agencia → NO necesita cita académica. Si quieres respaldarlo, puedes remitir a (Flamenco Agency, s. f.). Los nombres propios de instituciones (ferias, redes, teatros) tampoco llevan paréntesis APA.",
     "Flamenco Agency (s. f.) — opcional"),

    ("Intro",
     "Ballet Nacional de España (fecha)",
     "«…como el Ballet Nacional de España, creado en 1978 (Ballet Nacional de España, s. f.)…»  [corrige «años 1980» → 1978]",
     "Ballet Nacional de España – INAEM (s. f.)"),

    ("Intro",
     "Joaquín Cortés (compañía en los años 1990)",
     "«…o compañías privadas como la de Joaquín Cortés en los años 1990 (Danza.es, s. f.).»",
     "Danza.es – INAEM (s. f.)"),

    ("II.2",
     "Rasgos fonéticos del andaluz (elisión, aspiración de -s, apócope) — DEFINICIÓN lingüística",
     "«…el habla sevillana presenta varios rasgos fonéticos propios del dialecto andaluz (Narbona et al., 2011): la elisión consonántica, la aspiración de la -s…»",
     "Narbona et al. (2011). El español hablado en Andalucía"),

    ("III.2",
     "Noción de «centro de compra» (buying center) — TEORÍA",
     "«Suele intervenir lo que Webster y Wind (1972) denominaron el «centro de compra» (buying center): un conjunto de actores con funciones distintas.»  [sustituye «la literatura del sector»]",
     "Webster & Wind (1972)"),

    ("III.2",
     "El espectáculo como servicio intangible — TEORÍA",
     "«Como señalan Kotler y Keller (2012), un servicio es intangible: el cliente no puede «probarlo» antes de comprarlo.»",
     "Kotler & Keller (2012). Dirección de marketing"),

    ("III.2",
     "Estrategias push / pull — TEORÍA",
     "«…se proponen de forma activa en las campañas (estrategia push), … (estrategia pull) (Kotler & Keller, 2012).»",
     "Kotler & Keller (2012)"),

    ("III.3",
     "Vicente Escudero, triunfo en París (1920)",
     "«Ya en 1920, el bailaor Vicente Escudero triunfó en París (Archivo Vicente Escudero, s. f.) y demostró que el flamenco tenía una verdadera dignidad artística.»",
     "Archivo Vicente Escudero (s. f.)"),
]

t = doc.add_table(rows=1, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = 'Table Grid'
header_row(t, ["Apartado", "Elemento (no es tuyo)", "Cómo citarlo en el texto", "Fuente (→ sección 4)"])
for ap, el, cita, fu in rows:
    cells = t.add_row().cells
    set_cell(cells[0], ap, bold=True, size=9.5)
    set_cell(cells[1], el, size=9.5)
    set_cell(cells[2], cita, size=9.5)
    set_cell(cells[3], fu, size=9, color=GREY)

# largeurs
widths = [Cm(1.6), Cm(4.6), Cm(7.4), Cm(4.0)]
for row in t.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w

# ============================== 3. CORRECTIONS ==============================
h1("3. ⚠ Tres datos a corregir (son inexactos)")
note("Ces trois chiffres sont faux ou mal attribués dans la version actuelle — le jury peut le vérifier facilement.")

c = doc.add_table(rows=1, cols=3)
c.alignment = WD_TABLE_ALIGNMENT.CENTER
c.style = 'Table Grid'
header_row(c, ["Lo que dice ahora (✗)", "Problema", "Cómo escribirlo (✓)"])
corr = [
    ("«el sector cultural y turístico representan ≈ 12 % del PIB (Ministerio de Cultura y Deporte, 2023)»",
     "El 12 % corresponde al TURISMO, no a la cultura (≈ 2,4 %); y ese dato no es del Ministerio de Cultura.",
     "«El turismo representa aproximadamente el 12 % del PIB español (INE, 2023), mientras que el sector cultural aporta en torno al 2,4 % (Ministerio de Cultura, 2024).»"),
    ("«España ocupa entre el 5.º y el 8.º lugar mundial en congresos (ICCA, 2023)»",
     "Confusión país/ciudades: como PAÍS España es 3.ª; el 5.º/8.º son las CIUDADES (Barcelona/Madrid).",
     "«España es el tercer país del mundo en número de congresos internacionales, y Barcelona y Madrid figuran entre las diez ciudades más activas (ICCA, 2024).»"),
    ("«según Turespaña, más del 60 % de los turistas… incluyen una actividad cultural»",
     "Cifra no confirmada / probablemente obsoleta. El dato oficial reciente ≈ 19 % por motivo cultural.",
     "«Según Turespaña, cerca de uno de cada cinco viajes de turistas internacionales tiene una motivación cultural (Turespaña, 2023).»"),
]
for ahora, prob, ok in corr:
    cells = c.add_row().cells
    set_cell(cells[0], ahora, size=9.5)
    set_cell(cells[1], prob, size=9.5, color=GREY)
    set_cell(cells[2], ok, size=9.5)
cwidths = [Cm(5.5), Cm(5.0), Cm(7.1)]
for row in c.rows:
    for i, w in enumerate(cwidths):
        row.cells[i].width = w

note("Dato de Sevilla VERIFICADO: 3.018.060 viajeros alojados en hoteles de la ciudad en 2023 (INE, "
     "Encuesta de Ocupación Hotelera; recogido por el IECA). Por tanto «más de 3 millones» es correcto.")

# ============================== 4. TABLA 1 (GLOSARIO) ==============================
h1("4. Tabla 1 — Glosario de flamenco: ¿de dónde vienen las definiciones?")
note("⚠ Important : ne PAS laisser « elaboración propia » seul. Ces définitions s'appuient sur des "
     "dictionnaires — il faut les créditer. Les 16 termes ont été vérifiés : 13 figurent dans la RAE "
     "(palo, cantaor, bailaor, compás, cante jondo, zapateado, jaleo, mantón, castañuelas, duende, tablao, "
     "gira, bolo) ; 3 ne sont PAS dans la RAE et viennent d'un glossaire flamenco (tocaor, cuadro flamenco, "
     "bata de cola). Toutes les définitions sont conformes à ces sources.")

p = doc.add_paragraph()
p.add_run("Nota à mettre sous la Tabla 1 (remplace « Elaboración propia a partir de la experiencia… ») :").bold = True
box = doc.add_paragraph()
box.paragraph_format.left_indent = Cm(0.5)
br = box.add_run("Nota. Elaboración propia a partir del Diccionario de la lengua española "
                 "(Real Academia Española, s. f.) y del Glosario de términos del Centro Andaluz de "
                 "Flamenco (Junta de Andalucía, s. f.).")
br.italic = True
br.font.size = Pt(10.5)
br.font.color.rgb = NAVY

# ============================== 5. FUENTES ==============================
doc.add_page_break()
h1("5. Fuentes")
note("À copier-coller telle quelle à la fin du rapport (ordre alphabétique, sangría francesa, formato APA 7).")

fuentes = [
    "Archivo Vicente Escudero. (s. f.). Biografía. https://vicenteescudero.org/biografia/",
    "Ballet Nacional de España. (s. f.). Historia. Instituto Nacional de las Artes Escénicas y de la Música (INAEM), Ministerio de Cultura. https://balletnacional.inaem.gob.es",
    "Centro Andaluz de Flamenco. (s. f.). Glosario de términos. Junta de Andalucía. https://www.centroandaluzdeflamenco.es",
    "Danza.es. (s. f.). Biografías de danza española: Vicente Escudero y Joaquín Cortés. INAEM, Ministerio de Cultura. https://www.danza.es",
    "Flamenco Agency. (s. f.). Sitio web oficial. https://flamencoagency.com",
    "Instituto de Estadística y Cartografía de Andalucía. (2024). Encuesta de Coyuntura Turística de Andalucía (ECTA). Junta de Andalucía. https://www.juntadeandalucia.es/institutodeestadisticaycartografia",
    "Instituto Nacional de Estadística. (2023). Cuenta Satélite del Turismo de España. INE. https://www.ine.es",
    "International Congress and Convention Association. (2024). ICCA Statistics Report 2023: Country and city rankings. ICCA.",
    "Junta de Andalucía. (s. f.). El flamenco: patrimonio cultural inmaterial. Instituto Andaluz del Flamenco. https://www.juntadeandalucia.es/cultura/flamenco",
    "Kotler, P., & Keller, K. L. (2012). Dirección de marketing (14.ª ed.). Pearson Educación.",
    "Mailchimp. (s. f.). About open and click rates. https://mailchimp.com/help/about-open-and-click-rates/",
    "Ministerio de Cultura. (2024). Anuario de estadísticas culturales 2024. Gobierno de España. https://www.cultura.gob.es",
    "Ministerio de Cultura. (2024). Cuenta Satélite de la Cultura en España: avance de resultados 2020-2023. Gobierno de España.",
    "Narbona Jiménez, A., Cano Aguilar, R., & Morillo-Velarde Pérez, R. (2011). El español hablado en Andalucía. Universidad de Sevilla.",
    "Real Academia Española. (s. f.). Diccionario de la lengua española (edición del tricentenario, actualización 2023). https://dle.rae.es",
    "Spain Convention Bureau. (2023). Turismo de reuniones en España. Federación Española de Municipios y Provincias. https://scb.es",
    "Turespaña. (2023). Turismo cultural: perfil del turista internacional. Gobierno de España.",
    "UNESCO. (2010). El flamenco. Lista Representativa del Patrimonio Cultural Inmaterial de la Humanidad. https://ich.unesco.org/es/RL/el-flamenco-00363",
    "Webster, F. E., & Wind, Y. (1972). A general model for understanding organizational buying behavior. Journal of Marketing, 36(2), 12-19.",
]
for f in fuentes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)  # sangría francesa
    p.paragraph_format.space_after = Pt(6)
    p.add_run(f).font.size = Pt(10.5)

doc.save("Fuentes y citas - informe de practicas.docx")
print("Documento generado correctamente.")
